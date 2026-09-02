"""Convert Deep Sky Notes .docx articles to HTML pages."""

import re
import sys
import zipfile
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

SITE_ROOT = Path(__file__).resolve().parent.parent
ARTICLES_SRC = SITE_ROOT.parent / "Articles"
OUTPUT_DIR = SITE_ROOT / "articles"

MONTH_NAMES = {
    1: 'January', 2: 'February', 3: 'March', 4: 'April',
    5: 'May', 6: 'June', 7: 'July', 8: 'August',
    9: 'September', 10: 'October', 11: 'November', 12: 'December'
}

# Where each article appeared in Universe. The Published section was folded
# into Articles, so the PDFs now hang off the article they belong to.
# article slug -> (pdf stem in published/, volume, issue number, issue month)
ISSUE_MAP = {
    '2025-04': ('2025-06', 74, 6, 'June 2025'),
    '2025-05': ('2025-07', 74, 7, 'July 2025'),
    '2025-06': ('2025-08', 74, 8, 'August 2025'),
    '2025-08': ('2025-10', 74, 10, 'October 2025'),
    '2025-09': ('2025-11', 74, 11, 'November 2025'),
    '2026-01': ('2026-03', 75, 3, 'March 2026'),
    '2026-03': ('2026-05', 75, 5, 'May 2026'),
    '2026-04': ('2026-06', 75, 6, 'June 2026'),
    '2026-07': ('2026-08', 75, 8, 'August 2026'),
    '2026-08': ('2026-09', 75, 9, 'September 2026'),
}


def issue_credit(slug: str) -> str:
    """The `Published in Universe ...` line for an article's meta block."""
    entry = ISSUE_MAP.get(slug)
    if not entry:
        return 'Published in <em>Universe</em>, Astronomical Society of NSW'
    _, vol, iss, when = entry
    return (f'Published in <em>Universe</em> Vol {vol} #{iss:02d}, {when}, '
            f'Astronomical Society of NSW')


def issue_downloads(slug: str, published_dir: Path) -> str:
    """Download buttons for the published PDFs, or '' if none are on disk."""
    entry = ISSUE_MAP.get(slug)
    if not entry:
        return ''
    stem, vol, iss, _ = entry
    extract = published_dir / f'{stem}.pdf'
    if not extract.exists():
        return ''
    mb = extract.stat().st_size / (1024 * 1024)
    # Extract only. published/*-full.pdf is gitignored, so the complete issues
    # exist locally but never reach GitHub Pages -- linking one gives every
    # visitor a 404. (Also the author's standing rule: host the column extract,
    # not the whole issue.)
    button = (f'          <a href="../published/{extract.name}" class="download-btn" '
              f'download>&#8681; Download as published (PDF &middot; {mb:.1f} MB)</a>')
    return ('        <div class="article-downloads">\n'
            + button + '\n'
            + '        </div>\n')


def issue_card_link(slug: str) -> str:
    """The small `As published` line shown on an article-index card."""
    entry = ISSUE_MAP.get(slug)
    if not entry:
        return ''
    stem, vol, iss, _ = entry
    return (f'\n            <a class="card-pdf-link" href="../published/{stem}.pdf" '
            f'download>&#8681; As published &mdash; <em>Universe</em> Vol {vol} #{iss:02d}</a>')


def find_docx(folder: Path) -> Path | None:
    candidates = list(folder.glob('*.docx'))
    candidates = [c for c in candidates if not c.name.startswith('~') and 'Claude' not in c.name]
    if not candidates:
        return None
    for c in candidates:
        if 'Wanderings' in c.name or 'wanderings' in c.name:
            return c
    return candidates[0]


def find_separate_images(folder: Path) -> dict[int, Path]:
    images: dict[int, Path] = {}
    for f in sorted(folder.iterdir()):
        if f.name.startswith('~') or f.name.startswith('.'):
            continue
        if f.suffix.lower() not in ('.jpg', '.jpeg', '.png', '.heic'):
            continue
        name = f.stem
        m = re.match(r'(?:Figure|Image[_\s]?)[\s_]*(\d+)', name, re.IGNORECASE)
        if m:
            images[int(m.group(1))] = f
            continue
        m = re.match(r'IMG_(\d+)', name)
        if m:
            seq = int(m.group(1))
            idx = len(images) + 1
            images[idx] = f
    return images


def extract_embedded_images(docx_path: Path, out_dir: Path) -> dict[int, Path]:
    images: dict[int, Path] = {}
    out_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path) as z:
        media_files = sorted([n for n in z.namelist() if n.startswith('word/media/')])
        for i, name in enumerate(media_files, 1):
            ext = Path(name).suffix.lower()
            if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.emf', '.wmf'):
                continue
            if ext in ('.emf', '.wmf'):
                continue
            out_path = out_dir / f"figure-{i:02d}{ext}"
            out_path.write_bytes(z.read(name))
            images[i] = out_path
    return images


def parse_folder_date(folder_name: str) -> tuple[int, int]:
    m = re.match(r'(\d{4})_(\d{2})', folder_name)
    if m:
        return int(m.group(1)), int(m.group(2))
    return 2025, 1


def run_to_html(run) -> str:
    text = run.text or ''
    if not text:
        return ''
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    if run.bold:
        text = f'<strong>{text}</strong>'
    if run.italic:
        text = f'<em>{text}</em>'
    if run.font and run.font.superscript:
        text = f'<sup>{text}</sup>'
    if run.font and run.font.subscript:
        text = f'<sub>{text}</sub>'
    return text


def paragraph_to_html(para, figure_images: dict[int, str], figure_counter: list[int]) -> str:
    text_content = para.text.strip()

    is_figure_caption = bool(re.match(r'Figure\s+\d+', text_content, re.IGNORECASE))
    is_heading = para.style and para.style.name and 'Heading' in para.style.name

    blips = para._element.findall(
        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    has_image = len(blips) > 0

    if not text_content and not has_image:
        return ''

    if is_heading:
        level = 2
        m = re.search(r'(\d)', para.style.name)
        if m:
            level = min(int(m.group(1)) + 1, 4)
        inner = ''.join(run_to_html(r) for r in para.runs)
        return f'<h{level}>{inner}</h{level}>'

    if is_figure_caption:
        m = re.match(r'(Figure\s+(\d+))', text_content, re.IGNORECASE)
        if m:
            fig_num = int(m.group(2))
            img_path = figure_images.get(fig_num, '')
            caption_html = ''.join(run_to_html(r) for r in para.runs)
            if img_path:
                return f'''<figure>
  <img src="{img_path}" alt="{text_content}" loading="lazy">
  <figcaption>{caption_html}</figcaption>
</figure>'''
            else:
                return f'<p class="text-secondary"><em>{caption_html}</em></p>'

    figures_html = ''
    if has_image:
        for _ in blips:
            figure_counter[0] += 1
            fig_num = figure_counter[0]
            img_path = figure_images.get(fig_num, '')
            if img_path:
                figures_html += f'''<figure>
  <img src="{img_path}" alt="Figure {fig_num}" loading="lazy">
</figure>
'''
    if not text_content:
        return figures_html

    inner = ''.join(run_to_html(r) for r in para.runs)
    if not inner.strip():
        return figures_html if figures_html else ''

    alignment = para.alignment
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return figures_html + f'<p style="text-align:center">{inner}</p>'

    return figures_html + f'<p>{inner}</p>'


def convert_article(folder: Path, output_dir: Path, year: int, month: int,
                    prev_article: str | None = None, next_article: str | None = None) -> dict | None:
    docx_path = find_docx(folder)
    if not docx_path:
        print(f"  No .docx found in {folder}")
        return None

    separate_imgs = find_separate_images(folder)
    if separate_imgs:
        from scripts.optimize_images import optimize_image
        img_out = output_dir / 'images' / f'{year}-{month:02d}'
        img_out.mkdir(parents=True, exist_ok=True)
        figure_images: dict[int, str] = {}
        for num, src_path in sorted(separate_imgs.items()):
            dst_name = f'figure-{num:02d}.jpg'
            dst_path = img_out / dst_name
            optimize_image(src_path, dst_path, 1000)
            figure_images[num] = f'images/{year}-{month:02d}/{dst_name}'
    else:
        img_out = output_dir / 'images' / f'{year}-{month:02d}'
        embedded = extract_embedded_images(docx_path, img_out)
        figure_images = {}
        for num, src_path in sorted(embedded.items()):
            from scripts.optimize_images import optimize_image
            dst_name = f'figure-{num:02d}.jpg'
            dst_path = img_out / dst_name
            optimize_image(src_path, dst_path, 1000)
            figure_images[num] = f'images/{year}-{month:02d}/{dst_name}'

    doc = docx.Document(str(docx_path))

    figure_counter = [0]
    body_parts = []
    for para in doc.paragraphs:
        html = paragraph_to_html(para, figure_images, figure_counter)
        if html:
            body_parts.append(html)

    body_html = '\n'.join(body_parts)
    month_name = MONTH_NAMES.get(month, str(month))
    title = f'Deep Sky Notes &mdash; {month_name} {year}'
    slug = f'{year}-{month:02d}'

    first_img = figure_images.get(1, '')

    credit = issue_credit(slug)
    downloads = issue_downloads(slug, SITE_ROOT / 'published')

    nav_prev = ''
    nav_next = ''
    if prev_article:
        nav_prev = f'<a href="{prev_article}.html">&larr; Previous</a>'
    if next_article:
        nav_next = f'<a href="{next_article}.html">Next &rarr;</a>'

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <meta name="robots" content="noai, noimageai">
  <title>{title}</title>
  <meta name="description" content="Deep Sky Notes observing report for {month_name} {year}">
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="../css/style.css">
</head>
<body>

  <nav class="site-nav">
    <div class="nav-inner">
      <a href="../index.html" class="nav-brand">Deep Sky Notes</a>
      <button class="nav-toggle" aria-label="Toggle navigation" onclick="document.querySelector('.nav-links').classList.toggle('open')">&#9776;</button>
      <ul class="nav-links">
        <li><a href="index.html" class="active">Articles</a></li>
        <li class="nav-dropdown">
          <a href="../catalogue/index.html">Observing Notes &#9662;</a>
          <div class="nav-dropdown-menu">
            <a href="../catalogue/index.html">Master Table</a>
            <a href="../catalogue/constellations/index.html">Browse by Constellation</a>
          </div>
        </li>
        <li><a href="../observing-lists/index.html">Observing Lists</a></li>
        <li><a href="../links.html">Links</a></li>
        <li><a href="../about.html">About</a></li>
      </ul>
    </div>
  </nav>

  <main class="page-content">
    <div class="container">
      <div class="breadcrumbs">
        <a href="index.html">Articles</a><span class="sep">&rsaquo;</span>
        {month_name} {year}
      </div>

      <div class="article-header">
        <h1 class="mt-0">Deep Sky Notes &mdash; {month_name} {year}</h1>
        <div class="article-meta">
          <time datetime="{year}-{month:02d}">{month_name} {year}</time> &middot;
          {credit}
        </div>
{downloads}      </div>

      <div class="article-body">
        {body_html}
      </div>

      <div class="article-nav">
        {nav_prev or '<span></span>'}
        {nav_next or '<span></span>'}
      </div>
    </div>
  </main>

  <footer class="site-footer">
    <div class="container">
      <p>Deep Sky Notes &mdash; Alessandro Spina</p>

      <p class="text-secondary">Finder charts produced with <a href="https://stellarium.org" target="_blank" rel="noopener">Stellarium</a>. Survey images from the <a href="https://archive.stsci.edu/cgi-bin/dss_form" target="_blank" rel="noopener">Digitized Sky Survey</a> (STScI/NASA). Object data queried from the <a href="https://simbad.u-strasbg.fr/" target="_blank" rel="noopener">SIMBAD</a> database (CDS, Strasbourg).</p>
      <p class="text-secondary">&copy; 2026 Alessandro Spina. All rights reserved.</p>
    </div>
  </footer>

</body>
</html>'''

    out_file = output_dir / f'{slug}.html'
    out_file.write_text(html, encoding='utf-8')
    print(f"  Generated {out_file.name} ({len(figure_images)} images)")

    excerpt = ''
    for para in doc.paragraphs:
        t = para.text.strip()
        if t and not t.startswith('Figure') and len(t) > 50:
            excerpt = t[:150] + ('…' if len(t) > 150 else '')
            break

    return {
        'slug': slug,
        'year': year,
        'month': month,
        'monthName': month_name,
        'title': f'Deep Sky Notes — {month_name} {year}',
        'excerpt': excerpt,
        'thumbnail': first_img,
    }


def generate_article_index(articles: list[dict], output_dir: Path):
    cards = []
    for a in articles:
        thumb = ''
        if a['thumbnail']:
            thumb = f'<img class="article-card-thumb" src="{a["thumbnail"]}" alt="" loading="lazy">'
        cards.append(f'''
        <div class="article-card">
          {thumb}
          <div class="article-card-body">
            <h3><a href="{a['slug']}.html">{a['title']}</a></h3>
            <div class="date">{a['monthName']} {a['year']}</div>
            <p class="excerpt">{a['excerpt']}</p>{issue_card_link(a['slug'])}
          </div>
        </div>''')

    html = f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <meta name="robots" content="noai, noimageai">
  <title>Articles &mdash; Deep Sky Notes</title>
  <meta name="description" content="Monthly observing reports from Wiruna dark-sky site">
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="../css/style.css">
</head>
<body>

  <nav class="site-nav">
    <div class="nav-inner">
      <a href="../index.html" class="nav-brand">Deep Sky Notes</a>
      <button class="nav-toggle" aria-label="Toggle navigation" onclick="document.querySelector('.nav-links').classList.toggle('open')">&#9776;</button>
      <ul class="nav-links">
        <li><a href="index.html" class="active">Articles</a></li>
        <li class="nav-dropdown">
          <a href="../catalogue/index.html">Observing Notes &#9662;</a>
          <div class="nav-dropdown-menu">
            <a href="../catalogue/index.html">Master Table</a>
            <a href="../catalogue/constellations/index.html">Browse by Constellation</a>
          </div>
        </li>
        <li><a href="../links.html">Links</a></li>
        <li><a href="../about.html">About</a></li>
      </ul>
    </div>
  </nav>

  <main class="page-content">
    <div class="container">
      <h1>Articles</h1>
      <p class="text-secondary">Monthly observing reports from new-moon weekends at Wiruna, written as the <strong>Wiruna Wanderings</strong> column for <em>Universe</em>, the journal of the <a href="https://www.asnsw.com" target="_blank" rel="noopener">Astronomical Society of New South Wales</a>. Each report links the issue it appeared in, as a PDF.</p>

      <div class="article-list">{''.join(cards)}
      </div>
    </div>
  </main>

  <footer class="site-footer">
    <div class="container">
      <p>Deep Sky Notes &mdash; Alessandro Spina</p>

      <p class="text-secondary">Finder charts produced with <a href="https://stellarium.org" target="_blank" rel="noopener">Stellarium</a>. Survey images from the <a href="https://archive.stsci.edu/cgi-bin/dss_form" target="_blank" rel="noopener">Digitized Sky Survey</a> (STScI/NASA). Object data queried from the <a href="https://simbad.u-strasbg.fr/" target="_blank" rel="noopener">SIMBAD</a> database (CDS, Strasbourg).</p>
      <p class="text-secondary">&copy; 2026 Alessandro Spina. All rights reserved.</p>
    </div>
  </footer>

</body>
</html>'''

    (output_dir / 'index.html').write_text(html, encoding='utf-8')
    print(f"Generated article index with {len(articles)} entries")


def main():
    src = ARTICLES_SRC
    out = OUTPUT_DIR
    if len(sys.argv) > 1:
        src = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])

    out.mkdir(parents=True, exist_ok=True)

    folders = sorted([d for d in src.iterdir() if d.is_dir() and re.match(r'\d{4}_\d{2}', d.name)])
    slugs = []
    articles_meta = []

    for folder in folders:
        year, month = parse_folder_date(folder.name)
        slugs.append(f'{year}-{month:02d}')

    for i, folder in enumerate(folders):
        year, month = parse_folder_date(folder.name)
        prev_slug = slugs[i - 1] if i > 0 else None
        next_slug = slugs[i + 1] if i < len(slugs) - 1 else None
        print(f"Converting {folder.name}...")
        meta = convert_article(folder, out, year, month, prev_slug, next_slug)
        if meta:
            articles_meta.append(meta)

    articles_meta.reverse()
    generate_article_index(articles_meta, out)


if __name__ == '__main__':
    main()
